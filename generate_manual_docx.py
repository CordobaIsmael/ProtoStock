import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_manual():
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de colores
    COLOR_PRIMARY = RGBColor(225, 29, 72)     # Rose-600
    COLOR_SECONDARY = RGBColor(15, 23, 42)    # Slate-900
    COLOR_MUTED = RGBColor(100, 116, 139)     # Slate-500

    # Título Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("📖 MANUAL OFICIAL E INSTRUCTIVO DE USO")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("ProtoStock POS - Sistema de Gestión Comercial")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph() # Espacio

    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)

    # 1. Visión General
    add_h1("1. Visión General & Compatibilidad")
    p = doc.add_paragraph()
    p.add_run("ProtoStock es una solución integral diseñada para fiambrerías, almacenes, kioscos y mini-mercados. Permite administrar de forma ágil las ventas de mostrador, fraccionamiento por peso, control de caja y vencimientos.").font.size = Pt(11)
    
    add_bullet("Multidispositivo: ", "Compatible con PC, tablets (Android / iPad) y teléfonos móviles.")
    add_bullet("Periféricos Integrados: ", "Soporte nativo para lectoras de código de barras USB/Bluetooth, balanzas y tickoteras térmicas (58mm y 80mm).")
    add_bullet("Modo PWA Pantalla Completa: ", "Instalable como aplicación sin bordes ni pestañas.")

    # 2. Roles y Seguridad
    add_h1("2. Roles de Usuario & Seguridad (RBAC)")
    add_bullet("👑 Administrador (ADMIN): ", "Acceso total, cambio de contraseñas de todos los usuarios, retiros de caja, hitos de ventas y creación de personal.")
    add_bullet("👔 Encargado (ENCARGADO): ", "Supervisión de catálogo, precios, lotes, proveedores y compras. Lista de usuarios en modo lectura.")
    add_bullet("💵 Cajero (CAJERO): ", "Acceso enfocado exclusivamente a POS (Ventas) y Apertura/Cierre de turno. Bloqueado para ver ganancias.")

    # 3. Flujo Completo del Administrador
    add_h1("3. Flujo de Trabajo del Administrador (Admin User Flow)")
    p = doc.add_paragraph()
    p.add_run("A continuación se describe el paso a paso secuencial que realiza el Administrador en su jornada diaria de gestión:").font.size = Pt(11)

    add_h2("Paso 1: Inicio de Sesión & Autenticación")
    add_bullet("Ingreso: ", "Accede a la pantalla de Login (/login) y selecciona la cuenta @admin con su contraseña o PIN.")
    add_bullet("Redirección: ", "El sistema verifica las credenciales y lo redirige automáticamente al Dashboard principal.")

    add_h2("Paso 2: Revisión del Estado General (Dashboard)")
    add_bullet("Lectura de Métricas: ", "Visualiza el total facturado del día, los productos en nivel crítico de stock y los lotes próximos a vencer.")

    add_h2("Paso 3: Apertura de Turno de Caja")
    add_bullet("Apertura: ", "Ingresa a Turnos & Caja (/caja) para iniciar el turno indicando el monto en efectivo de cambio inicial (ej: $10.000).")

    add_h2("Paso 4: Operación en Mostrador (POS)")
    add_bullet("Modo Tablet: ", "En caso de operar desde una tablet, activa el 'Modo Tablet: Teclado Virtual OFF'.")
    add_bullet("Escaneo y Pesaje: ", "Escanea los productos con lector de barras o selecciona fiambres por gramos (100g, 250g, 500g).")
    add_bullet("Cobro e Impresión: ", "Selecciona el método de pago (Efectivo, Débito, Crédito, Mercado Pago), calcula el vuelto con billetes rápidos e imprime el Ticket Térmico.")

    add_h2("Paso 5: Gestión de Inventario, Precios y Lotes")
    add_bullet("Actualización de Precios: ", "Desde Productos & Precios (/productos), modifica precios de lista o importa nuevos catálogos desde Excel.")
    add_bullet("Carga de Compras: ", "Desde Compras & Lotes (/compras), registra nuevos lotes de mercadería ingresando proveedor y fecha de expiración.")
    add_bullet("Bajas de Stock: ", "Registra mermas o productos vencidos descontándolos de la base de datos.")

    add_h2("Paso 6: Auditoría de Seguridad & Gestión de Personal")
    add_bullet("Usuarios y Claves: ", "Desde Usuarios & Permisos (/usuarios), crea nuevos usuarios o modifica las contraseñas/PIN de cajeros y encargados.")
    add_bullet("Auditoría de Retiros: ", "Supervisa las salidas de efectivo registradas por empleados para pago a proveedores.")

    add_h2("Paso 7: Cierre de Jornada & Análisis de Reportes")
    add_bullet("Arqueo Final: ", "Ingresa a Turnos & Caja (/caja), realiza el recuento físico de billetes y cierra el turno registrando si hubo sobrante o faltante.")
    add_bullet("Reportes de Ganancia: ", "En Reportes & Métricas (/reportes), analiza la ganancia bruta acumulada y el desglose de ingresos por método de pago.")

    # 4. Módulo POS
    add_h1("4. Módulo POS: Punto de Venta & Cobro")
    add_bullet("Buscador F2 & Escáner: ", "Búsqueda instantánea por código de barras, SKU o nombre.")
    add_bullet("Modo Tablet (Teclado Virtual OFF): ", "Desactiva el despliegue del teclado de Android/iPadOS para que no tape la pantalla mientras se escanea.")
    add_bullet("Venta Fraccionada por Peso: ", "Accesos directos para 100g, 150g, 200g, 250g, 300g, 500g, 750g y 1kg con cálculo automático.")
    add_bullet("Calculadora de Vuelto & Billetes Rápidos: ", "Botones de billetes ($1.000, $2.000, $5.000, $10.000, $20.000, $50.000) para un cálculo instantáneo del vuelto.")

    # 5. Impresión Térmica
    add_h1("5. Módulo Impresión de Ticket Térmico")
    add_bullet("Formatos 58mm y 80mm: ", "Adaptable a tickoteras portátiles mini y de escritorio.")
    add_bullet("Comprobantes: ", "Soporta modo Control Interno (No Fiscal) y Vista AFIP (Factura C).")
    add_bullet("Impresión Limpia: ", "Genera un documento aislado sin encabezados o URLs del navegador.")

    # 6. Productos y Excel
    add_h1("6. Productos, Precios & Importación desde Excel")
    add_bullet("Importación Masiva: ", "Carga de más de 500 productos en 30 segundos importando un archivo CSV/Excel.")
    add_bullet("Exportación de Catálogo: ", "Descarga la base de datos de productos a CSV en 1 clic.")
    add_bullet("Control de Mermas: ", "Baja de stock indicando motivos (Vencido, Devolución, Rotura, Consumo Interno).")

    # 7. Caja Diaria
    add_h1("7. Turnos & Caja Diaria")
    add_bullet("Apertura de Turno: ", "Ingreso del dinero en efectivo de cambio inicial.")
    add_bullet("Egresos y Retiros: ", "Registro de salidas de dinero para proveedores con concepto descriptivo.")
    add_bullet("Arqueo Automático: ", "Cálculo de dinero esperado vs contado con reporte de sobrante o faltante.")

    # 8. Compras y Lotes
    add_h1("8. Compras, Lotes & Proveedores Frecuentes")
    add_bullet("Proveedores Frecuentes: ", "Gestión de datos de contacto de distribuidores.")
    add_bullet("Carga de Lotes con Vencimiento: ", "Control de fecha de expiración y stock por lote.")

    # 9. Reportes
    add_h1("9. Reportes & Métricas de Ganancia")
    add_bullet("Desglose Métodos de Pago: ", "Gráficos independientes para Efectivo, Débito, Crédito y Mercado Pago.")
    add_bullet("Ganancia Bruta: ", "Métrica de utilidad neta acumulada del negocio.")

    # 10. Notificaciones Table
    add_h1("10. Centro de Notificaciones Inteligentes por Rol")
    p = doc.add_paragraph("El sistema despliega alertas en vivo según el rol de la persona iniciada:")

    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Tipo de Alerta", "Roles Destinatarios", "Descripción"]
    data = [
        ["⚠️ Stock Mínimo", "Admin, Encargado, Cajero", "Aviso de productos por agotarse (stock <= min)."],
        ["⏳ Lotes a Vencer", "Admin, Encargado, Cajero", "Aviso de lotes a vencer en los próximos 30 días."],
        ["⏰ Recordatorio Arqueo", "Admin, Encargado, Cajero", "Alerta si la caja lleva más de 8 hs abierta."],
        ["💸 Retiros de Caja", "Solo ADMINISTRADOR", "Registro de salidas de dinero para proveedores."],
        ["🏆 Hito Facturación", "Solo ADMINISTRADOR", "Aviso al alcanzar metas de ventas del día."]
    ]

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            if r_idx % 2 == 1:
                shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shd)

    # Guardar archivo Word
    output_path = "d:/ProtoStock/MANUAL_USUARIO_PROTOSTOCK.docx"
    try:
        doc.save(output_path)
        print(f"Documento Word actualizado en: {output_path}")
    except PermissionError:
        fallback_path = "d:/ProtoStock/MANUAL_USUARIO_PROTOSTOCK_V2.docx"
        doc.save(fallback_path)
        print(f"Documento Word actualizado en: {fallback_path}")

if __name__ == "__main__":
    create_manual()
